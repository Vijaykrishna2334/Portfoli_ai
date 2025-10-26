# document_assets.py
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.colors import HexColor

import docx # Already installed
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from schema import ResumeProfile # To ensure we use the structured data

# --- A. PDF GENERATION FUNCTION (using ReportLab) ---

def generate_pdf(profile: ResumeProfile) -> BytesIO:
    """Generates a professional PDF resume using ReportLab."""
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            leftMargin=Inches(0.5), rightMargin=Inches(0.5),
                            topMargin=Inches(0.5), bottomMargin=Inches(0.5))
    
    styles = getSampleStyleSheet()
    Story = []

    # Custom Styles
    style_h1 = styles['Heading1']
    style_h1.fontSize = 20
    style_h1.textColor = HexColor('#007bff')
    style_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    style_h2 = styles['Heading2']
    style_h2.fontSize = 14
    style_h2.spaceBefore = 10
    style_h2.spaceAfter = 5
    style_h2.textColor = HexColor('#333333')
    
    style_body = styles['Normal']
    style_body.fontSize = 10
    style_body.spaceAfter = 5

    # 1. Name and Contact (Header)
    Story.append(Paragraph(profile.name, style_h1))
    Story.append(Paragraph(profile.email, styles['Italic']))
    Story.append(Spacer(1, 0.2*Inches))

    # 2. Summary
    Story.append(Paragraph("SUMMARY", style_h2))
    Story.append(Paragraph(profile.summary, style_body))
    Story.append(Spacer(1, 0.1*Inches))
    
    # 3. Skills
    Story.append(Paragraph("KEY SKILLS", style_h2))
    Story.append(Paragraph(", ".join(profile.skills), style_body))
    Story.append(Spacer(1, 0.1*Inches))

    # 4. Experience
    Story.append(Paragraph("EXPERIENCE", style_h2))
    
    for exp in profile.experience:
        Story.append(Paragraph(f"<b>{exp.title}</b> at {exp.company}", styles['Heading3']))
        Story.append(Paragraph(exp.years, styles['Italic']))
        
        # Use ListFlowable for professional bullet points
        bullet_items = [ListItem(Paragraph(exp.summary, style_body))]
        Story.append(ListFlowable(bullet_items, bulletType='bullet', start='bullet'))
        Story.append(Spacer(1, 0.1*Inches))

    doc.build(Story)
    buffer.seek(0)
    return buffer

# --- B. DOCX GENERATION FUNCTION (using python-docx) ---

def generate_docx(profile: ResumeProfile) -> BytesIO:
    """Generates an ATS-ready DOCX resume."""
    
    document = Document()

    # 1. Name and Contact (Header)
    p = document.add_paragraph()
    runner = p.add_run(profile.name)
    runner.font.size = Pt(18)
    runner.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph(profile.email).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph().add_run('—'*50).bold = True

    # 2. Summary
    document.add_heading('Summary', level=2)
    document.add_paragraph(profile.summary)
    
    # 3. Skills
    document.add_heading('Key Skills', level=2)
    document.add_paragraph(", ".join(profile.skills))

    # 4. Experience
    document.add_heading('Work Experience', level=2)
    
    for exp in profile.experience:
        document.add_paragraph(f"{exp.title} - {exp.company}", style='List Bullet')
        document.add_paragraph(f"Dates: {exp.years}").italic = True
        document.add_paragraph(exp.summary)
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer