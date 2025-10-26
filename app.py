import streamlit as st
import os
import json
from dotenv import load_dotenv 
from typing import List
from io import BytesIO

# Import the Google GenAI SDK and necessary schema utilities
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch


# --- SCHEMA DEFINITIONS ---
class WorkExperience(BaseModel):
    title: str = Field(description="The user's job title at this company.")
    company: str = Field(description="The name of the company.")
    years: str = Field(description="The start and end date/year range (e.g., '2020 - 2023').")
    summary: str = Field(description="A 2-3 sentence summary of responsibilities and achievements.")

class ResumeProfile(BaseModel):
    name: str = Field(description="The user's full name.")
    email: str = Field(description="The user's professional email address.")
    summary: str = Field(description="A concise, professional 3-sentence summary of the user's career goals and experience.")
    skills: List[str] = Field(description="A list of 8 to 12 key hard skills (e.g., Python, SQL, React).")
    experience: List[WorkExperience] = Field(description="A list of all work experiences.")

class OptimizationReport(BaseModel):
    match_score: int = Field(description="A confidence score from 0 to 100...")
    keyword_gaps: List[str] = Field(description="A list of 3-5 critical skills...")
    suggestions: List[str] = Field(description="A list of 3 actionable, specific suggestions...")

class InterviewSettings(BaseModel):
    role: str = Field(description="The target job role...")
    type: str = Field(description="The type of questions...")
    status: str = Field(default="ready", description="The current status...")
    history: List[str] = Field(default=[], description="List of all turns...")
# --- END SCHEMA DEFINITIONS ---


# --- 0. Setup and Configuration ---
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY") 

# --- CRITICAL FIX: CACHE THE GEMINI CLIENT ---
@st.cache_resource
def get_gemini_client():
    if not API_KEY:
        st.warning("🚨 GEMINI_API_KEY not found. Some AI features will be disabled.")
        return None
    try:
        return genai.Client(api_key=API_KEY)
    except Exception as e:
        st.error(f"Error initializing Gemini client: {e}")
        return None

client = get_gemini_client()

if not client:
    st.stop() 


# --- DOCUMENT HANDLING AND GENERATION FUNCTIONS ---

def read_pdf(file):
    reader = PdfReader(file)
    text = "".join(page.extract_text() for page in reader.pages)
    return text

def read_docx(file):
    document = docx.Document(file)
    text = "".join(para.text + "\n" for para in document.paragraphs)
    return text

def generate_pdf_v2(profile: ResumeProfile) -> BytesIO:
    """Generates a professional PDF resume using ReportLab (V2 Feature)."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.5*inch, rightMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    Story = []
    styles.add(ParagraphStyle(name='HeadingName', fontName='Helvetica-Bold', fontSize=20, alignment=1, textColor=HexColor('#007bff')))
    styles.add(ParagraphStyle(name='HeadingSection', fontName='Helvetica-Bold', fontSize=14, spaceBefore=12, spaceAfter=4, textColor=HexColor('#333333')))
    styles.add(ParagraphStyle(name='NormalSmall', fontName='Helvetica', fontSize=10, leading=12))
    styles.add(ParagraphStyle(name='BulletStyle', fontName='Helvetica', fontSize=10, leftIndent=0.25*inch, bulletIndent=-0.25*inch, spaceBefore=0, spaceAfter=0)) 

    Story.append(Paragraph(profile.name.upper(), styles['HeadingName']))
    Story.append(Paragraph(profile.email, styles['Italic']))
    Story.append(Spacer(1, 0.2*inch))
    Story.append(Paragraph("SUMMARY", styles['HeadingSection']))
    Story.append(Paragraph(profile.summary.strip(), styles['NormalSmall'])) 
    Story.append(Spacer(1, 0.1*inch))
    Story.append(Paragraph("KEY SKILLS", styles['HeadingSection']))
    Story.append(Paragraph(f"Skills: {', '.join(profile.skills)}", styles['NormalSmall']))
    Story.append(Spacer(1, 0.1*inch))
    Story.append(Paragraph("WORK EXPERIENCE", styles['HeadingSection']))
    
    for exp in profile.experience:
        Story.append(Paragraph(f"<b>{exp.title}</b> at {exp.company}", styles['NormalSmall']))
        Story.append(Paragraph(exp.years, styles['Italic']))
        clean_summary = exp.summary.replace('·', '').strip()
        bullet_items = [ListItem(Paragraph(line.strip(), styles['NormalSmall']), bulletText='\u2022') for line in clean_summary.split('\n') if line.strip()]
        if bullet_items:
            Story.append(ListFlowable(bullet_items, bulletType='bullet', start='bullet'))
        Story.append(Spacer(1, 0.1*inch))
    doc.build(Story)
    buffer.seek(0)
    return buffer

def generate_docx_v2(profile: ResumeProfile) -> BytesIO:
    """Generates an ATS-ready DOCX resume with proper structural styling (V2 Feature)."""
    document = Document()
    p = document.add_paragraph()
    runner = p.add_run(profile.name.upper())
    runner.font.size = Pt(18)
    runner.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(profile.email).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph().add_run('—'*50).font.size = Pt(8) 
    document.add_heading('SUMMARY', level=2)
    document.add_paragraph(profile.summary.strip()) 
    document.add_heading('KEY SKILLS', level=2)
    document.add_paragraph(", ".join(profile.skills))
    document.add_heading('EXPERIENCE', level=2)
    
    for exp in profile.experience:
        p_title = document.add_paragraph()
        p_title.add_run(f"{exp.title}").bold = True
        p_title.add_run(f" | {exp.company} ({exp.years})").italic = True
        for line in exp.summary.split('\n'):
             clean_line = line.strip()
             if clean_line:
                 document.add_paragraph(clean_line, style='List Bullet')
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- CORE AI FUNCTIONS ---
def parse_resume_to_json(resume_text: str) -> ResumeProfile:
    if not client: st.warning("AI client not available."); return None
    
    prompt = ("You are an expert career data analyst. Your job is to extract all relevant "
        "professional information from the user's resume text and structure it "
        "according to the provided JSON schema...")

    try:
        response = client.models.generate_content(model="gemini-2.5-flash", contents=[prompt + f"---RESUME TEXT---\n{resume_text}"], config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=ResumeProfile))
        return ResumeProfile(**json.loads(response.text))
    except Exception as e:
        st.error(f"AI Parsing Error: {e}"); return None

def generate_cover_letter(profile_json: str, job_description: str, tone: str) -> str:
    if not client: return "Error: Gemini client not initialized."
    
    prompt = f"""You are an expert career coach and professional writer. Your task is to write a highly persuasive... (omitted prompt body)"""
    try:
        response = client.models.generate_content(model="gemini-2.5-flash", contents=[prompt])
        return response.text
    except Exception as e:
        st.error(f"Cover Letter API Error: {e}"); return None

def generate_optimization_report(profile_json: str, job_description: str) -> dict:
    if not client: return None
    
    prompt = f"""You are an expert Applicant Tracking System (ATS)... (omitted prompt body)"""
    try:
        response = client.models.generate_content(model="gemini-2.5-flash", contents=[prompt], config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=OptimizationReport))
        return json.loads(response.text)
    except Exception as e:
        st.error(f"Optimization API Error: {e}"); return None

def generate_final_feedback(history: list, role: str, type: str, profile_json: str) -> str:
    if not client: return "Error: Gemini client not initialized."
    
    transcript = "\n".join([f"**{m['role'].upper()}:** {m['text']}" for m in history])
    prompt = f"""You are a professional Interview Performance Analyst. Analyze the following interview transcript... (omitted prompt body)"""

    try:
        response = client.models.generate_content(model="gemini-2.5-flash", contents=[prompt])
        return response.text
    except Exception as e:
        st.error(f"Final Feedback API Error: {e}"); return "Failed to generate final feedback report."
# --- END CORE AI FUNCTIONS ---


# --- DASHBOARD (Main Application Content) ---

st.set_page_config(
    page_title="PortfolioAI Builder (V1 - Open Access)",
    layout="wide",
    initial_sidebar_state="expanded"
)
st.title("🤖 Instant PortfolioAI Builder (Open Access MVP)")
st.caption("Welcome! This version allows direct access to all AI features without a login.")

# --- 1. Data Input Section (Updated to handle files) ---
st.header("1. Input Your Data")
st.caption("You can either **upload a file** or **paste the text** below.")

uploaded_file = st.file_uploader("Upload PDF or Word Document (.docx)", type=["pdf", "docx"])
user_data = st.text_area("OR, Paste Your Resume/CV Content Here:", height=300, placeholder="Example:\nName: Alice Smith...")

resume_text_to_parse = ""
parsing_source = "Pasted Text"
if uploaded_file is not None:
    file_extension = uploaded_file.name.split('.')[-1]
    parsing_source = f"Uploaded {file_extension.upper()} File"
    with st.spinner(f"Reading text from {uploaded_file.name}..."):
        if file_extension == "pdf": resume_text_to_parse = read_pdf(uploaded_file)
        elif file_extension == "docx": resume_text_to_parse = read_docx(uploaded_file)
elif user_data:
    resume_text_to_parse = user_data

# --- 2. Template Selection ---
st.header("2. Choose Your Design")
template_choice = st.selectbox("Select a Portfolio Template:", options=["Minimalist (Template 1)", "Technical (Template 2)", "Creative (Template 3)"])

# --- 3. Asset Generation and Display (Parsing Trigger) ---
st.header("3. Your Instant Portfolio")

if 'parsed_profile' not in st.session_state: st.session_state.parsed_profile = None

if st.button("Generate & Publish Portfolio 🚀"):
    if not resume_text_to_parse: st.error("Please provide content either by uploading a file or pasting text.")
    elif not client: st.error("Cannot run AI features. Gemini client failed to initialize.")
    else:
        with st.spinner(f"🧠 Analyzing and Structuring Data from {parsing_source} with Gemini 2.5 Flash..."):
            profile: ResumeProfile = parse_resume_to_json(resume_text_to_parse) 
            st.session_state.parsed_profile = profile
            
    if st.session_state.parsed_profile:
        st.success(f"🎉 Success! Data parsed and portfolio generated using **{template_choice}** template.")
        
        st.subheader("✅ Profile Data Stored (Raw JSON Review)")
        st.json(st.session_state.parsed_profile.model_dump()) 

        st.subheader("Live Preview: Key Data Points (For Review)")
        profile = st.session_state.parsed_profile
        st.markdown(f"**Name:** {profile.name}")
        st.markdown(f"**Email:** {profile.email}")
        st.markdown("**Summary:**")
        st.write(profile.summary) 
        st.markdown("**Key Skills:**")
        st.write(", ".join(profile.skills))
        st.markdown("**Work Experience:**")
        for exp in profile.experience:
            st.markdown(f"**- {exp.title}** at {exp.company} ({exp.years})")
            st.write(f"  * {exp.summary}")


        st.subheader("Actionable Assets")
        
        # --- Build Content for Placeholder Assets ---
        profile = st.session_state.parsed_profile
        txt_resume_content = (f"--- ATS-READY TEXT PROFILE ---\n\nName: {profile.name}\nEmail: {profile.email}\nSummary:\n{profile.summary}\n\nSkills: {', '.join(profile.skills)}\n\n")
        
        # --- FINAL DOWNLOAD BUTTONS (V2-Ready) ---
        
        st.info("Download buttons below provide static placeholders for V1 stability.")
        
        col_txt, col_word, col_pdf = st.columns(3)
        
        # 1. TXT Button (Fully functional)
        with col_txt:
            st.download_button("Download Profile as Text File (.txt)", data=txt_resume_content.encode('utf-8'), file_name="ATS_Resume.txt", key="profile_download_txt", mime="text/plain")
        
        # 2. DOCX Button (V2 Ready)
        with col_word:
            docx_bytes = generate_docx_v2(profile).getvalue()
            st.download_button(
                "Download Profile as Word Document (.docx)", 
                data=docx_bytes, 
                file_name="ATS_Resume_V2.docx",
                key="profile_download_docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # 3. PDF Button (V2 Ready)
        with col_pdf:
            pdf_bytes = generate_pdf_v2(profile).getvalue()
            st.download_button(
                "Download Profile as PDF (.pdf)", 
                data=pdf_bytes, 
                file_name="ATS_Resume_V2.pdf",
                key="profile_download_pdf",
                mime="application/pdf"
            )
        
    # The error message for failure is handled inside parse_resume_to_json.

# --- 4. AI Cover Letter Writer Section ---
st.divider()
st.header("📝 AI Cover Letter Writer")
st.caption("Generate a hyper-tailored cover letter by pasting the Job Description.")

job_description = st.text_area("Paste the Target Job Description Here:", height=250, placeholder="Paste the full text of the job description you are applying for. The AI will match your profile to it.")
tone_choice = st.selectbox("Select Desired Tone:", options=["Professional & Formal", "Friendly & Enthusiastic", "Technical & Detail-Oriented"], key="cover_letter_tone_select")

if st.button("Generate Cover Letter 🖋️", key="generate_cover_letter_button"): 
    if not st.session_state.parsed_profile: st.error("Please run the 'Generate & Publish Portfolio' step first to parse your resume data.")
    elif not job_description: st.error("Please paste the Job Description to tailor the letter.")
    else:
        with st.spinner("✍️ Writing Persuasive, Tailored Cover Letter with Gemini..."):
            cover_letter_text = generate_cover_letter(st.session_state.parsed_profile.model_dump_json(), job_description, tone_choice)
            if cover_letter_text:
                st.success("Cover letter complete!")
                st.text_area("Generated Cover Letter", cover_letter_text, height=400)
                st.download_button("Download Cover Letter (TXT)", cover_letter_text, file_name="Cover_Letter.txt", key="cover_letter_download_button")
            else:
                st.error("Cover letter generation failed.")


# --- 5. AI Optimizer & Scorer Section ---
st.divider()
st.header("5. 🧠 AI Optimizer & Job Matcher")
st.caption("Maximize your interview chances: Compare your profile against a target job description.")

target_job_description = st.text_area("Paste the Target Job Description for Analysis:", height=250, placeholder="Paste the full text of the job description (JD) here. The AI will analyze your strengths and gaps.")

if st.button("Analyze Profile & Get Score ⭐", key="analyze_profile_button"): 
    if not st.session_state.parsed_profile: st.error("Please run the 'Generate & Publish Portfolio' step first to parse your resume data.")
    elif not target_job_description: st.error("Please paste the Job Description to perform the analysis.")
    else:
        with st.spinner("🔬 Performing Keyword Gap Analysis and Optimization..."):
            optimization_result = generate_optimization_report(st.session_state.parsed_profile.model_dump_json(), target_job_description)
            if optimization_result:
                st.success("Analysis complete!")
                score = optimization_result.get("match_score", 0)
                gaps = optimization_result.get("keyword_gaps", ["N/A"])
                suggestions = optimization_result.get("suggestions", ["N/A"])
                st.metric(label="Match Confidence Score", value=f"{score}%", delta="High score = high ATS compatibility")
                st.subheader("Key Improvement Areas (Gaps)")
                st.warning("⚠️ The following keywords/skills from the JD are missing or underrepresented in your profile:")
                st.markdown("\n".join([f"- **{gap}**" for gap in gaps]))
                st.subheader("Optimization Suggestions")
                st.info("💡 Use these suggestions to update your profile and maximize your score:")
                st.markdown("\n".join([f"- {s}" for s in suggestions]))
            else:
                st.error("Optimization analysis failed.")


# --- 6. AI Mock Interviewer Section ---
st.divider()
st.header("6. 🎙️ AI Mock Interviewer")
st.caption("Get realistic practice and real-time coaching for your target role.")

if "interview_settings" not in st.session_state: st.session_state.interview_settings = InterviewSettings(role="", type="", status="ready") 
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "chat_session" not in st.session_state: st.session_state.chat_session = None
if "final_report_display" not in st.session_state: st.session_state.final_report_display = None

col1, col2 = st.columns(2)
with col1: target_role = st.text_input("Target Role (e.g., AI Engineer):", value=st.session_state.interview_settings.role, key="interview_role_input")
with col2: question_type = st.selectbox("Question Type:", options=["Behavioral (STAR)", "Technical", "System Design"], index=["Behavioral (STAR)", "Technical", "System Design"].index(st.session_state.interview_settings.type) if st.session_state.interview_settings.type in ["Behavioral (STAR)", "Technical", "System Design"] else 0, key="interview_type_select")

if st.session_state.interview_settings.status != "in_progress":
    if st.session_state.final_report_display:
        st.subheader("Final Feedback Report")
        st.markdown(st.session_state.final_report_display)
        if st.button("Start New Interview Session", key="start_new_session_button"):
            st.session_state.final_report_display = None 
            st.session_state.interview_settings = InterviewSettings(role="", type="", status="ready") 
            st.rerun() 
        
    else: 
        if st.button("Start Interview Session ▶️", key="start_interview_button"): 
            if st.session_state.parsed_profile is None: st.error("Please generate your profile in Section 3 first.")
            elif not target_role: st.error("Please enter a Target Role.")
            else:
                system_instruction = f"""You are a professional, rigorous technical interviewer. You are interviewing the candidate for the role of '{target_role}' focused on '{question_type}' questions. The candidate's profile data is: {st.session_state.parsed_profile.model_dump_json()} Rules: 1. You must start by welcoming the candidate and asking the first question ONLY. DO NOT INCLUDE ANY FEEDBACK YET. 2. After this initial welcome and question, provide brief, constructive feedback on the candidate's previous response and ask the next question in the sequence. 3. Do not rush or give too much information at once. Maintain a professional tone. 4. The interview should last for 5 questions total."""
                st.session_state.chat_session = client.chats.create(model="gemini-2.5-flash", history=[types.Content(role="user", parts=[types.Part.from_text(text=system_instruction)])])
                with st.spinner("Initializing interview and generating first question..."):
                    initial_response = st.session_state.chat_session.send_message("Please start the interview now.")
                st.session_state.chat_history = []
                st.session_state.chat_history.append({"role": "model", "text": initial_response.text})
                st.session_state.interview_settings = InterviewSettings(role=target_role, type=question_type, status="in_progress")
                st.rerun() 
            
else: 
    st.subheader(f"Interview in Progress: {st.session_state.interview_settings.role} ({st.session_state.interview_settings.type})")
    for message in st.session_state.chat_history:
        role = "assistant" if message['role'] == "model" else "user"
        with st.chat_message(role): st.write(message['text'])
        
    if prompt := st.chat_input("Your response...", key="chat_input_key"): 
        st.session_state.chat_history.append({"role": "user", "text": prompt})
        with st.chat_message("user"): st.write(prompt)
        with st.chat_message("assistant"):
            with st.spinner("AI Interviewer is analyzing and preparing the next question..."):
                response = st.session_state.chat_session.send_message(prompt)
            st.session_state.chat_history.append({"role": "model", "text": response.text})
            st.write(response.text)

    if st.button("Finish Interview & Get Full Feedback 🛑", key="finish_interview_button"): 
        st.session_state.interview_settings.status = "finished"
        st.subheader("Final Feedback Report")
        with st.spinner("Analyzing transcript and generating final performance report..."):
            final_report = generate_final_feedback(history=st.session_state.chat_history, role=st.session_state.interview_settings.role, type=st.session_state.interview_settings.type, profile_json=st.session_state.parsed_profile.model_dump_json())
            st.session_state.final_report_display = final_report
        st.session_state.interview_settings = InterviewSettings(role="", type="", status="finished")
        st.rerun() 

# --- 7. Job-Opening Alert Engine (V1 Placeholder) ---
st.divider()
st.header("7. 📧 Job-Opening Alert Engine")
st.caption("Targeted job discovery for higher application efficiency.")
st.warning("⚠️ This feature is a high-priority V2 roadmap item.")
st.info("In V1, we are collecting user interest and basic preferences.")
if st.session_state.parsed_profile:
    st.markdown(f"**Based on your profile ({st.session_state.parsed_profile.summary[:30]}...),**")
    job_keywords = st.text_input("Enter Target Job Titles/Keywords (e.g., 'MLOps, NLP, Junior Data Scientist'):", placeholder="Enter your keywords", key="job_alert_keywords")
    st.select_slider("Set Alert Frequency:", options=['Daily', 'Weekly', 'Bi-Weekly', 'Monthly'], value='Weekly', key="alert_frequency_slider")
    if st.button("Save Job Alert Preferences", key="save_alerts_button"): st.success(f"Preferences saved for future launch! We will notify you when personalized job alerts are live.")
else: st.info("Generate your portfolio in Section 3 to unlock job alert preference collection!")